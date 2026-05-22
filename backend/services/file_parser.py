import os
import pdfplumber
from docx import Document


def parse_file(file_path: str) -> str:
    ext = file_path.lower().split(".")[-1]
    if ext == "pdf":
        return _parse_pdf(file_path)
    elif ext == "docx":
        return _parse_docx(file_path)
    elif ext in ["jpg", "jpeg", "png"]:
        return _parse_image(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}")


def parse_first_page(file_path: str) -> str:
    ext = file_path.lower().split(".")[-1]
    if ext == "pdf":
        try:
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) > 0:
                    return pdf.pages[0].extract_text() or ""
                return ""
        except Exception as e:
            print(f"PDF First Page Parse Error: {e}")
            return ""
    elif ext == "docx":
        try:
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs[:20]])
        except Exception as e:
            print(f"DOCX First Page Parse Error: {e}")
            return ""
    return parse_file(file_path)


def get_chunks(file_path: str) -> list:
    """파일을 페이지/단락 단위 청크 리스트로 반환"""
    ext = file_path.lower().split(".")[-1]
    if ext == "pdf":
        return _pdf_chunks(file_path)
    elif ext == "docx":
        return _docx_chunks(file_path)
    else:
        text = parse_file(file_path)
        return [text] if text.strip() else []


def _pdf_chunks(file_path: str) -> list:
    chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                parts = []
                try:
                    found_tables = page.find_tables()
                    text_page = page
                    for t in found_tables:
                        text_page = text_page.outside_bbox(t.bbox)
                    text = text_page.extract_text()
                    if text:
                        parts.append(text)
                    for t in found_tables:
                        rows = []
                        for row in t.extract():
                            cells = [str(c or '').strip().replace('\n', ' ') for c in row]
                            rows.append(' | '.join(cells))
                        if rows:
                            parts.append('[표]\n' + '\n'.join(rows))
                except Exception as e:
                    print(f"PDF page chunk error (fallback): {e}")
                    text = page.extract_text()
                    if text:
                        parts.append(text)
                if parts:
                    chunks.append('\n'.join(parts).strip())
    except Exception as e:
        print(f"PDF chunking error: {e}")
    return chunks


def _docx_chunks(file_path: str) -> list:
    chunks = []
    try:
        doc = Document(file_path)
        current = []
        for para in doc.paragraphs:
            if para.text.strip():
                current.append(para.text)
            if len("\n".join(current)) > 2000:
                chunks.append("\n".join(current))
                current = []
        if current:
            chunks.append("\n".join(current))
        for table in doc.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            chunks.append("[표]\n" + "\n".join(rows))
    except Exception as e:
        print(f"DOCX chunking error: {e}")
    return chunks


def _parse_pdf(file_path: str) -> str:
    with pdfplumber.open(file_path) as pdf:
        pages = []
        for page in pdf.pages:
            parts = []
            try:
                found_tables = page.find_tables()
                text_page = page
                for t in found_tables:
                    text_page = text_page.outside_bbox(t.bbox)
                text = text_page.extract_text()
                if text:
                    parts.append(text)
                for t in found_tables:
                    rows = []
                    for row in t.extract():
                        cells = [str(c or '').strip().replace('\n', ' ') for c in row]
                        rows.append(' | '.join(cells))
                    if rows:
                        parts.append('[표]\n' + '\n'.join(rows))
            except Exception as e:
                print(f"PDF page parse error (fallback): {e}")
                text = page.extract_text()
                if text:
                    parts.append(text)
            if parts:
                pages.append('\n'.join(parts))
        return '\n\n'.join(pages)


def _parse_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _parse_image(file_path: str) -> str:
    from services.gemini_client import call_gemini_vision
    return call_gemini_vision(
        file_path,
        "이 이미지에서 텍스트를 전부 추출해줘. 표나 항목 구조가 있으면 그대로 유지해줘."
    )
